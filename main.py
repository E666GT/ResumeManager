import docx
import pandas
import comtypes.client
from win32com import client as wc
import os
from docx.shared import Cm, Inches
import qtpy
from MainUI import Ui_Form
import time
import sys
from PyQt5.QtWidgets import QApplication , QMainWindow, QAction,QSizePolicy,QLabel,QPushButton,QMessageBox,QWidget
from PyQt5.QtCore import QThread,pyqtSignal
from PyQt5.QtGui import QTextCursor
import datetime
import shutil
import subprocess
import threading
today=datetime.date.today()
formatted_today=today.strftime('%y%m%d')
# sys.path.append("..")
"""
# Author:  Yiwen
# 
# Refs:
    # style:#https://python-docx.readthedocs.io/en/latest/user/styles-using.html
    # table https://python-docx.readthedocs.io/en/latest/api/table.html
"""


SOURCE_EXCEL_PATH = "AllInOneSource.xlsx"  # all in one source.xlsx
TEMPLETE_WORD_PATH = "style_templete.docx"  # convert source_excel <=> world.docx

SAVE_WORD_PATH = "CV_XiaoMing.docx"
SAVE_PDF_PATH = "CV_XiaoMing.pdf"
# SAVE_TXT = ""

# SAVE_WORD_EDIT_PATH = "output_EDIT.docx"
# SAVE_EXCEL_EDIT_PATH = "output_EDIT.xlsx"
# SAVE_TXT_LOC_EDIT = "output_EDIT.docx.loc.txt"

SAVE_WORD_EDIT_PATH = "EDIT.docx"
SAVE_EXCEL_EDIT_PATH = SOURCE_EXCEL_PATH #"EDIT.xlsx" # same as SOURCE_EXCEL_PATH, every time updated with EDIT.docx
SAVE_TXT_LOC_EDIT = "EDIT.docx.loc.txt"

TYPE = "MA"  # CS ME AI MA EDIT

# size
SIZE_TIME_WIDTH = 5  # cm
SIZE_TITLE_WIDTH = 13  # cm

class MainUI(QMainWindow,Ui_Form):
    def __init__(self,parent=None,op=None):
        super(MainUI, self).__init__(parent)
        self.setupUi(self)
        self.op=op

        self.label_block.setEnabled(False)
        self.label_block.setVisible(False)
        # self.excelpath_label.setText(op.source_excel)
        # self.templeteword_path_label.setText(op.templete_word)
        self.excelpath_lineEdit.setText(op.source_excel)
        self.excelpath_lineEdit.setEnabled(False)
        self.templete_word_lineEdit.setText(op.templete_word)
        self.templete_word_lineEdit.setEnabled(False)
        self.wordpath_lineEdit.setText(SAVE_WORD_EDIT_PATH)
        self.wordpath_lineEdit.setEnabled(False)

        self.cvtype_lineEdit.setText("ME")

        self.button_word2excel.clicked.connect(self.docx2excel)
        self.button_excel2word.clicked.connect(self.excel2word)
        self.button_excel2cv.clicked.connect(self.excel2cv)



        #update modified time
        self.info("获取修改时间...\n")
        self.QtThread_Update_mtime = QtThread_Update_mtime()
        self.QtThread_Update_mtime.sin_update.connect(self.update_mtime)
        self.QtThread_Update_mtime.ui_logic_update.connect(self.update_ui_logic)
        self.QtThread_Update_mtime.start()

        #init labels list
        self.info("初始化labels\n")
        self.listWidget_labels.addItems(self.op.get_label_finished_list())
        self.listWidget_labels.clicked.connect(self.choose_label)

        self.pushButton_openExcel.clicked.connect(self.open_excel)
        self.pushButton_openDoc.clicked.connect(self.open_doc)

        self.info("程序启动完毕！\n")
    def open_excel(self):
        # openthread=threading.Thread(target=self.open_file,args=("AllInOneSource.xlsx",))
        # openthread.start()

        self.label_block.setVisible(True)
        self.info("保存关闭打开的文件后，才可下一步操作。")
        self.open_file("AllInOneSource.xlsx")
        self.label_block.setVisible(False)
        pass
    def open_doc(self):

        # openthread=threading.Thread(target=self.open_file,args=("EDIT.docx",))
        # openthread.start()

        self.label_block.setVisible(True)
        self.info("保存关闭打开的文件后，才可下一步操作。")
        self.open_file("EDIT.docx")
        self.label_block.setVisible(False)
        pass
    def open_file(self,file):
        cmd = os.path.join(os.getcwd(),file)
        returned_value = subprocess.call(cmd, shell=True)
        print(returned_value)
    def info(self,str):
        # self.DebugBrowser.scroll

        self.DebugBrowser.insertPlainText(str+"\n")
        self.DebugBrowser.moveCursor(QTextCursor.End)
        # self.DebugBrowser.setText(text+"\n"+str)
    def choose_label(self):
        item = self.listWidget_labels.currentItem()
        print("    Clicked-"+item.text())
        self.cvtype_lineEdit.setText(item.text())
        # self.cvtype_label.setText(item.text())
    def update_ui_logic(self):

        if(self.checkBox_export_all_labels_cv.isChecked()):
            self.listWidget_labels.setEnabled(False)
        else:
            self.listWidget_labels.setEnabled(True)

    def update_mtime(self):
        self.label_allinone.setText(self.op.get_mtime(self.excelpath_lineEdit.text()))
        self.label_editdoc.setText(self.op.get_mtime(self.wordpath_lineEdit.text()))
        if(self.checkBox_export_all_labels_cv.isChecked()):
            self.listWidget_labels.setEnabled(False)
        else:
            self.listWidget_labels.setEnabled(True)
        self.listWidget_labels.clear()
        self.listWidget_labels.addItems(self.op.get_label_finished_list())
    def docx2excel(self):
        self.op.convert_docx_edit_2_excel(SAVE_WORD_EDIT=self.wordpath_lineEdit.text(),SAVE_EXCEL_EDIT=self.excelpath_lineEdit.text())
    def excel2word(self):
        self.op.convert_excel_2_docx_edit(excel_f=self.excelpath_lineEdit.text(),output_word_path=self.wordpath_lineEdit.text())
    def excel2cv(self):
        if self.checkBox_export_all_labels_cv.isChecked():
            labels_list=self.op.get_labels_list()
            langs=["C","E"]
            for label in labels_list:
                for lang in langs:
                    self.info("正在导出"+label+"_"+lang)
                    self.op.convert_excel_2_docx(self.excelpath_lineEdit.text(),
                                                 target_type=label,
                                                 output_word_name=SAVE_WORD_PATH,
                                                 pdf=True,
                                                 language=lang)

        else:
            try:
                self.op.set_templete(self.templete_word_lineEdit.text())
                self.op.paras["TYPE"]=self.cvtype_lineEdit.text()
                if(self.checkBox_LangChinese.isChecked()):
                    lang="C"
                else:
                    lang="E"
                self.op.convert_excel_2_docx(self.excelpath_lineEdit.text(),target_type=self.cvtype_lineEdit.text(),output_word_name=SAVE_WORD_PATH,pdf=self.check_ispdf.isChecked(),language=lang)
                self.DebugBrowser.setText("Finish!")
            except Exception as err:
                print(err)
    # #EXCEL -> EDIT_WORD
    # AF.convert_excel_2_docx_edit(excel_f=SOURCE_EXCEL_PATH,output_word_path=SAVE_WORD_EDIT_PATH)
    #
    # #EDIT_WORD -> EXCEL
    # # AF.convert_docx_edit_2_excel(SAVE_WORD_EDIT=SAVE_WORD_EDIT_PATH,SAVE_EXCEL_EDIT=SAVE_EXCEL_EDIT_PATH)
    #
    # #EXCEL -> CV
    # AF.convert_excel_2_docx(SAVE_EXCEL_EDIT_PATH,"ME",output_word_name=SAVE_WORD_PATH)


class QtThread_Update_mtime(QThread):
    sin_update = pyqtSignal(bool)

    ui_logic_update = pyqtSignal(bool)

    def __init__(self, parent=None):
        super(QtThread_Update_mtime, self).__init__(parent)
        # #设置工作状态与初始num数值
        # self.working = True
        # self.num = 0

    def __del__(self):
        # #线程状态改变与线程终止
        # self.working = False
        # self.wait()
        pass

    def run(self):
        loop=0
        while True:
            # print(loop)
            loop=loop+1
            time.sleep(0.01)
            if(loop%20==0):
                self.sin_update.emit(1)
            self.ui_logic_update.emit(1)
            if(loop==1000):
                loop=0
            # #获取文本
            # file_str = 'File index{0}'.format(self.num)
            # self.num += 1
            # 发射信号
            # 线程休眠2秒
            # self.sleep(2)

class AutoFormer(object):
    def __init__(self,SOURCE_EXCEL,TEMPLETE_WORD):
        self.source_excel = SOURCE_EXCEL
        self.blank_document = docx.Document(TEMPLETE_WORD)
        self.templete_word=TEMPLETE_WORD
        # self.transdict = {
        #     "WorkExp": "PROFESSIONAL EXPERIENCE",
        #     "Summary": "SUMMARY",
        #     "Education": "EDUCATION",
        #     "Projects": "RESEARCH EXPERIENCE",
        #     "Campus": "UNIVERSITY/PUBLIC ENGAGEMENTS",
        #     "Pub": "SELECTED PUBLICATION / CONFERENCE",
        #     "Awards": "AWARDS & SCHOLARSHIPS",
        #     "Skill": "TECHNICAL SKILLS/ADDITIONAL SKILLS",
        #     "Info": " "
        # }
        self.language="E"
        self.paras = {
            "TYPE": TYPE,
            "TIME_WIDTH": 4,  # CM,
            "TITLE_WIDTH": 16  # CM,

        }
        print(self.get_labels_list())
        pass
    @property
    def transdict(self):
        transdict={}
        if(self.language=="E"):
            if(self.paras["TYPE"]=="MA"):
                transdict = {
                    "WorkExp": "PROFESSIONAL EXPERIENCE",
                    "Summary": "SUMMARY",
                    "Education": "EDUCATION",
                    "Projects": "RESEARCH EXPERIENCE / COMPUS PROJECTS",
                    "Campus": "UNIVERSITY/PUBLIC ENGAGEMENTS",
                    "Pub": "SELECTED PUBLICATION / CONFERENCE",
                    "Awards": "AWARDS & SCHOLARSHIPS",
                    "Skill": "TECHNICAL SKILLS/ADDITIONAL SKILLS",
                    "Info": " "
                }
            else:
                transdict = {
                    "WorkExp": "PROFESSIONAL EXPERIENCE",
                    "Summary": "SUMMARY",
                    "Education": "EDUCATION",
                    "Projects": "RESEARCH EXPERIENCE",
                    "Campus": "UNIVERSITY/PUBLIC ENGAGEMENTS",
                    "Pub": "SELECTED PUBLICATION / CONFERENCE",
                    "Awards": "AWARDS & SCHOLARSHIPS",
                    "Skill": "TECHNICAL SKILLS/ADDITIONAL SKILLS",
                    "Info": " "
                }
        if(self.language=="C"):
            if (self.paras["TYPE"]=="MA"):
                transdict = {
                    "WorkExp": "工作经历",
                    "Summary": "概要",
                    "Education": "教育经历",
                    "Projects": "项目经历",
                    "Campus": "校园和社会经历",
                    "Pub": "发表",
                    "Awards": "奖学金与荣誉",
                    "Skill": "技能",
                    "Info": " "
                }
            else:
                transdict = {
                    "WorkExp": "工作经历",
                    "Summary": "概要",
                    "Education": "教育经历",
                    "Projects": "研究经历",
                    "Campus": "校园和社会经历",
                    "Pub": "发表",
                    "Awards": "奖学金与荣誉",
                    "Skill": "技能",
                    "Info": " "
                }

        return transdict
        # pass
    # ========================
    # Functions
    # ========================
    def convert_excel_2_docx(self, excel_f,target_type,output_word_name,pdf,language):
        print("EXCEL -> DOCX....")
        print("    EXCEL=",excel_f)
        print("    target_type=",target_type)
        # print("    output_word_name=",output_word_name)
        print("    language=",language)
        self.backup_allinonecsv()
        self.language=language
        self.paras["TYPE"]=target_type

        document = self.get_templete_document()
        # document.save("test.docx")
        excel_f = pandas.ExcelFile(excel_f)
        # print(excel_f.sheet_names)
        Include_Sheets = ["Info", "Summary", "Education", "WorkExp", "Projects",  "Pub","Campus","Awards","Skill"]
        for sheet_name in Include_Sheets:
            df = excel_f.parse(sheet_name=sheet_name)
            # print(df.shape)
            print("    convert_excel_2_docx - Add Module:", sheet_name)
            if 'Weight' in df.columns:
                df.sort_values(by="Weight", ascending=False, inplace=True)
            if sheet_name!="Info":
                document.add_paragraph("", style="正文 无项目")
            self.add_doc_module(sheet_name, df, document)

        folder = "output"
        save_word_name = output_word_name.split(".")
        save_word_name_ = os.path.join(folder,save_word_name[0] + "_" + self.paras["TYPE"] +"_"+str(language)+ "_"+str(formatted_today)+"." + "docx")


        document.save(save_word_name_)
        if(pdf):
            pdf_name=save_word_name[0] + "_" + self.paras["TYPE"]+"_"+str(language)+ "_"+str(formatted_today)+"." + "pdf"
            pdf_path=os.path.join(folder,pdf_name)

            self.convert_word_to_pdf(save_word_name_,pdf_path)
            print("    已导出pdf --",pdf_path)
        pass
        print("Finished!")

    def convert_excel_2_docx_edit(self, excel_f,output_word_path):

        print("EXCEL -> DOCX....")
        print("        Now Convert Excel file to Docx file")
        self.backup_allinonecsv()
        # document = self.blank_document
        # self.delete_paragraph(document.paragraphs[0])
        document = self.get_templete_document()
        excel_f = pandas.ExcelFile(excel_f)
        txt_f = open(SAVE_TXT_LOC_EDIT, "w+")
        # print(excel_f.sheet_names)
        # Include_Sheets = ["Info", "Summary", "Education", "WorkExp", "Projects", "Campus","Skill"]
        for sheet_name in excel_f.sheet_names:
            df = excel_f.parse(sheet_name=sheet_name)
            # print(df.shape)
            print("    Add Module:", sheet_name)
            # if 'Name' in df.columns:
            #     df.sort_values(by="Name",ascending=False,inplace=True)
            # document.add_paragraph("", style="正文 无项目")
            # document.add_paragraph(str(df.to_json()), style="正文 无项目")
            # self.add_doc_module(sheet_name, df, document)
            for dfrow in df.iterrows():
                row_index = dfrow[0]
                col_loop = 0
                for col_name in dfrow[1]._stat_axis._data:
                    Location_Sig = [sheet_name, col_name, row_index]
                    # document.add_paragraph(str(Location_Sig))
                    txt_f.write(str(Location_Sig) + ";")
                    if col_loop == 0:
                        document.add_paragraph(str(dfrow[1][col_name]), style="Heading 2")
                    else:
                        if(str(dfrow[1][col_name])=="nan"):
                            document.add_paragraph(str("请在此处输入"+col_name+":"))
                        else:
                            document.add_paragraph(str("请在此处输入"+col_name+":")+str(dfrow[1][col_name]))

                    col_loop += 1
                pass
        txt_f.close()
        document.save(output_word_path)
        # self.convert_word_to_pdf(SAVE_WORD,SAVE_PDF)
        pass
        print("Finished!")

    def convert_docx_edit_2_excel(self, SAVE_WORD_EDIT,SAVE_EXCEL_EDIT):
        print("DOCX -> EXCEL....")
        self.backup_allinonecsv()
        document = docx.Document(SAVE_WORD_EDIT)
        paragraphs=[]
        txt_f = open(SAVE_TXT_LOC_EDIT, "r")
        for x in txt_f:
            locate_sig_str = x

        locate_sig_ = locate_sig_str.split(";")
        locate_sig=[]
        for i in range(len(locate_sig_)):
            if locate_sig_[i] != "":
                locate_sig.append(eval(str(locate_sig_[i])))
                # print(locate_sig[i])
        txt_f.close()
        for para in document.paragraphs:
            paragraphs.append(para.text)

        print("    paragraphs,总段落,",len(paragraphs))
        print("    locate_sig,总长度,",len(locate_sig))

        lass_sheet_name=""
        # sheet_names=[]
        dicts={}
        for i in range(len(paragraphs)):
            p=paragraphs[i]
            sig=locate_sig[i]
            sheet_name=sig[0]
            col_name=sig[1]
            row_index=sig[2]

            if(not sheet_name in dicts.keys()):
                # sheet_names.append(sheet_name)
                dicts[sheet_name]={}
            if(not col_name in dicts[sheet_name].keys()):
                dicts[sheet_name][col_name]={}

            if(type(dicts[sheet_name][col_name])!=list):
                dicts[sheet_name][col_name] = []
            if(p==("请在此处输入"+col_name+":")):
                p="nan"
            if(p[:6]=="请在此处输入"):
                p=p.replace("请在此处输入"+col_name+":","")
            dicts[sheet_name][col_name].append(p)

            # dicts[sheet_name][col_name].append()
        with pandas.ExcelWriter(SAVE_EXCEL_EDIT) as writer:
            for sheet_name in dicts.keys():
                save_pd=pandas.DataFrame(dicts[sheet_name])
                save_pd.to_excel(writer,sheet_name=sheet_name,index=False)

        print("Finished!")

    def generate_form(self, template):
        pass

    def save_template(self):
        pass

    # ========================
    # Modules
    # ========================
    def add_doc_module(self, sheet_name, df_sheet, document):
        headname = self.transdict[sheet_name]
        if (sheet_name != "Info"):
            document.add_heading(headname, level=1)
        for dfrow in df_sheet.iterrows():
            if 'Labels' in dfrow[1]._stat_axis._data:
                labels = str(dfrow[1]["Labels"]).split(';')
                if (self.paras["TYPE"] in labels):
                    pass
                else:
                    continue
            if "Lang" in dfrow[1]._stat_axis._data:
                lang= str(dfrow[1]["Lang"])
                if (self.language == lang):
                    pass
                else:
                    continue
            if "Finished" in dfrow[1]._stat_axis._data:
                Finished = str(dfrow[1]["Finished"])

                if (Finished == "Y"):
                    pass
                else:
                    continue

            self.add_row_module(sheet_name, dfrow[1], document)
    def add_row_module(self, sheet_name, df_row, document):
        if (sheet_name == "WorkExp"):
            table = document.add_table(rows=3, cols=2)
            table.cell(1, 0).merge(table.cell(1, 1))
            table.cell(2, 0).merge(table.cell(2, 1))
            r0 = table.rows[0].cells
            # row_cells = table.add_row().cells #add row
            r0[0].text = str(df_row["Position"])
            r0[0].paragraphs[0].style = document.styles['Heading 2']
            r0[0].width = Cm(self.paras['TITLE_WIDTH'])
            r0[1].width = Cm(self.paras['TIME_WIDTH'])
            r0[1].text = str(df_row["Time"])
            r0[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            r0[1].paragraphs[0].style = document.styles['正文 无项目']

            r1 = table.rows[1].cells
            r1[0].text = str(df_row["Company"]) + "," + str(df_row["Location"])
            r1[0].paragraphs[0].style = document.styles['正文 无项目']

            r2 = table.rows[2].cells
            r2[0].text = str(df_row["Achievement_1"])
            other_achieves = [df_row["Achievement_2"], df_row["Achievement_3"], df_row["Achievement_4"]]
            for achive in other_achieves:
                if pandas.isna(achive) :
                    continue
                r2[0].add_paragraph(str(achive))
        if (sheet_name == "Projects"):
            table = document.add_table(rows=3, cols=2)
            table.cell(1, 0).merge(table.cell(1, 1))
            table.cell(2, 0).merge(table.cell(2, 1))
            r0 = table.rows[0].cells
            # row_cells = table.add_row().cells #add row
            r0[0].text = df_row["Name"]
            r0[0].paragraphs[0].style = document.styles['Heading 2']

            r0[0].width = Cm(self.paras['TITLE_WIDTH'])
            r0[1].width = Cm(self.paras['TIME_WIDTH'])

            r0[1].text = str(df_row["Time"])
            r0[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            r0[1].paragraphs[0].style = document.styles['正文 无项目']

            r1 = table.rows[1].cells
            r1[0].text = df_row["Position"]
            r1[0].paragraphs[0].style = document.styles['正文 无项目']
            r1[0].add_paragraph(str(df_row["Orgnization"]) + "," + str(df_row["Location"]), style='正文 无项目')

            r2 = table.rows[2].cells
            r2[0].text = str(df_row["Achievement_1"])
            other_achieves = [df_row["Achievement_2"], df_row["Achievement_3"], df_row["Achievement_4"]]
            for achive in other_achieves:
                if pandas.isna(achive):
                    continue
                r2[0].add_paragraph(str(achive))
        if (sheet_name == "Summary"):
            document.add_paragraph(str(df_row["Sentence"]), style="正文 无项目")
        if (sheet_name == "Education"):
            # print(df_row)
            # document.add_paragraph(str(df_row["School"]), style="正文 无项目")

            table = document.add_table(rows=2, cols=2)
            table.cell(1, 0).merge(table.cell(1, 1))
            # table.cell(2, 0).merge(table.cell(2, 1))
            r0 = table.rows[0].cells
            # row_cells = table.add_row().cells #add row
            r0[0].text = df_row["School"]
            r0[0].paragraphs[0].style = document.styles['Heading 2']
            r0[0].width = Cm(self.paras['TITLE_WIDTH'])
            r0[1].width = Cm(self.paras['TIME_WIDTH'])
            r0[1].text = df_row["Time"]
            r0[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            r0[1].paragraphs[0].style = document.styles['正文 无项目']

            r1 = table.rows[1].cells

            degree_str = str(df_row["Degree"])+" - "+str(df_row["Major"])
            check_list = ["Major GPA", "GPA"]
            for it in check_list:
                if (pandas.isna(df_row[it])):
                    continue
                degree_str += ("," + it + ":" + df_row[it])
            r1[0].text = degree_str
            r1[0].paragraphs[0].style = document.styles['正文 无项目']
        if (sheet_name == "Info"):
            table = document.add_table(rows=3, cols=2)
            table.cell(0, 0).merge(table.cell(0, 1))
            table.rows[0].cells[0].text = df_row['Name']
            table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            table.rows[0].cells[0].paragraphs[0].style = "抬头姓名"

            table.rows[1].cells[0].paragraphs[0].style = "正文 无项目"
            table.rows[1].cells[1].paragraphs[0].style = "正文 无项目"

            table.rows[2].cells[0].text = df_row['Mobile1'] + "/" + df_row["Mobile2"]
            table.rows[2].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            table.rows[2].cells[0].paragraphs[0].style = "正文 无项目"

            table.rows[2].cells[1].text = df_row['Email']
            table.rows[2].cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            table.rows[2].cells[1].paragraphs[0].style = "正文 无项目"
        if (sheet_name == "Campus"):
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = str(df_row["Position"]) + "," + str(df_row["Name"])
            table.rows[0].cells[0].paragraphs[0].style = document.styles['Heading 2']
            table.rows[0].cells[1].text = str(df_row["Time"])
            table.rows[0].cells[1].paragraphs[0].style = document.styles["正文 无项目"]
            table.rows[0].cells[0].width = Cm(self.paras['TITLE_WIDTH'])
            table.rows[0].cells[1].width = Cm(self.paras['TIME_WIDTH'])
            table.rows[0].cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        if (sheet_name == "Skill"):
            document.add_paragraph(str(df_row["Type"]) + "  -  " + str(df_row["Name"]), style="正文 无项目")
        if (sheet_name == "Awards"):
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = str(df_row["Name"]) + "," + str(df_row["Rank"])
            table.rows[0].cells[0].paragraphs[0].style = document.styles["正文 无项目"]
            table.rows[0].cells[1].text = str(df_row["Time"])
            table.rows[0].cells[1].paragraphs[0].style = document.styles["正文 无项目"]
            table.rows[0].cells[0].width = Cm(self.paras['TITLE_WIDTH'])
            table.rows[0].cells[1].width = Cm(self.paras['TIME_WIDTH'])
            table.rows[0].cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        if (sheet_name == "Pub"):
            cont=str(df_row["Authors"])+","+str(df_row["Title"])+"("+str(df_row["Status"])+")"+str(df_row["Conf"])
            document.add_paragraph(cont, style="正文 无项目")
    # ========================
    # Macros
    # ========================
    # def clear_doc(self,):
    def get_label_finished_dict(self):
        excel_f = pandas.ExcelFile(self.source_excel)
        sheetnames = excel_f.sheet_names
        labels_list = []
        finish_dict={}
        for sheetname in sheetnames:
            if(sheetname=="全局备注"):
                pass
            else:
                continue
            df_sheet = excel_f.parse(sheet_name=sheetname)
            # df_sheet.
            for dfrow in df_sheet.iterrows():
                t_label=str(dfrow[1]["labels"])
                e_finish=str(dfrow[1]["英文完成度"])
                c_finish=str(dfrow[1]["中文完成度"])
                finish_dict[t_label]={
                    "e_finish":eval(e_finish),
                    "c_finish":eval(c_finish)
                }

        return finish_dict
    def get_label_finished_list(self):
        finished_dict=self.get_label_finished_dict()
        candidates_labels=self.get_labels_list()
        res_labels=[]
        for label in finished_dict.keys():
            valid=finished_dict[label]["c_finish"]+finished_dict[label]["e_finish"]
            if(valid>1):
                res_labels.append(label)
        return res_labels
    def get_labels_list(self):
        excel_f=pandas.ExcelFile(self.source_excel)
        sheetnames=excel_f.sheet_names
        labels_list=[]
        for sheetname in sheetnames:
            df_sheet = excel_f.parse(sheet_name=sheetname)
            # df_sheet.
            for dfrow in df_sheet.iterrows():
                if 'Labels' in dfrow[1]._stat_axis._data:
                    labels = str(dfrow[1]["Labels"]).split(';')
                    for label in labels:
                        if(label in labels_list):
                            pass
                        else:
                            labels_list.append(label)

        try:
            labels_list.remove("")
        except:
            pass
        try:
            labels_list.remove("nan")
        except:
            pass
        return labels_list
    def backup_allinonecsv(self):
        timestamp=time.strftime("%Y%m%d_%H%M%S")
        name=self.source_excel.split('.')[0]

        shutil.copy2(self.source_excel,os.path.join("backup",name+"_"+timestamp+".xlsx"))
        print("    已经备份AllInOne.csv至backup文件夹")
    def set_templete(self,templete_docx_path):
        self.templete_word=templete_docx_path
    def save_doc(self):
        pass
    def get_templete_document(self):
        word_path=self.templete_word
        document=docx.Document(word_path)
        while(len(document.paragraphs)):
            self.delete_paragraph(document.paragraphs[0])
        # for i in range():
        #     self.delete_paragraph(document.paragraphs[i])
        # for paragraph in document.paragraphs:
        #     # print(paragraph.text)
        #     self.delete_paragraph(paragraph)
        return document
    def get_mtime(self,path):
        return str(time.ctime(os.path.getmtime(path)))
    # ========================
    # Statics
    # ========================
    def test(self):
        pass

    @staticmethod
    def apply_style(para, style, document):
        para.style = document.styles[style]
    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    # @staticmethod
    # def convert_word_to_pdf(inputFile, outputFile):
    #     ''' the following lines that are commented out are items that others shared with me they used when
    #     running loops to stop some exceptions and errors, but I have not had to use them yet (knock on wood) '''
    #     word = comtypes.client.CreateObject('Word.Application')
    #     # word.visible = True
    #     # time.sleep(3)
    #     doc = word.Documents.Open(inputFile)
    #     doc.SaveAs(outputFile, FileFormat=17)
    #     doc.close()
    #     # word.visible = False
    #     word.Quit()
    @staticmethod
    def convert_word_to_pdf(docx_path, pdf_path):
        # DocxPaths = get_docx(input_Docxs)
        # print('DocxPaths', DocxPaths)
        word = wc.Dispatch('Word.Application')
        word.Visible = 0
        # current_path=os.getcwd()
        docx_path = os.path.join(os.getcwd(), docx_path)
        pdf_path = os.path.join(os.getcwd(), pdf_path)
        # pdf_name = pdf_path + '/' + docx_path.split('/')[-1][:-5] + '.pdf'
        try:
            doc = word.Documents.Open(docx_path)
            doc.SaveAs(pdf_path, 17)  # 直接保存为PDF文件
            # doc.Close()
        except Exception as err:
            print(err)
            print('%s 转换失败' % docx_path)
        word.Quit()


if __name__ == "__main__":

    AF = AutoFormer(SOURCE_EXCEL_PATH,TEMPLETE_WORD_PATH)
    # AF.convert_excel_2_docx(AF.source_excel)
    # AF.convert_excel_2_docx_edit(AF.source_excel)

    # #EXCEL -> EDIT_WORD
    # AF.convert_excel_2_docx_edit(excel_f=SOURCE_EXCEL_PATH,output_word_path=SAVE_WORD_EDIT_PATH)
    #
    # EDIT_WORD -> EXCEL
    # AF.convert_docx_edit_2_excel(SAVE_WORD_EDIT=SAVE_WORD_EDIT_PATH,SAVE_EXCEL_EDIT=SAVE_EXCEL_EDIT_PATH)
    #
    # EXCEL -> CV/
    # AF.convert_excel_2_docx(excel_f=SAVE_EXCEL_EDIT_PATH,target_type="FIN",output_word_name=SAVE_WORD_PATH,pdf=False,language="C")


    app = QApplication(sys.argv)
    win = MainUI(op=AF)
    win.show()
    sys.exit(app.exec_())

    print(AF.get_label_finished_dict())